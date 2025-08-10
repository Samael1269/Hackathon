import streamlit as st
import hashlib
import os
from jamaibase import JamAI, protocol as p
from docx import Document
from docx.shared import Pt
from io import BytesIO
from tempfile import NamedTemporaryFile
from PIL import Image
import base64
# ---------------- PAGE CONFIG (only once, at the very top) ----------------
st.set_page_config(
    page_title="NurseJoy.ai",
    page_icon="📝",
    layout="wide"
)

# ---------------- UTILS ----------------
def make_hashes(password):
    return hashlib.sha256(str.encode(password)).hexdigest()

def check_hashes(password, hashed_text):
    return make_hashes(password) == hashed_text

def load_users(file_path="users.txt"):
    users = {}
    if os.path.exists(file_path):
        with open(file_path, "r") as f:
            for line in f:
                line = line.strip()
                if line:
                    parts = line.split(",", 2)
                    if len(parts) == 3:
                        username, hashed, role = parts
                        users[username] = {"password": hashed, "role": role}
    return users

def save_user(username, hashed_password, role, file_path="users.txt"):
    with open(file_path, "a") as f:
        f.write(f"{username},{hashed_password},{role}\n")

# ---------------- SESSION STATE ----------------
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "username" not in st.session_state:
    st.session_state.username = ""
if "role" not in st.session_state:
    st.session_state.role = ""

# ---------------- LOAD USERS ----------------
users_db = load_users()

# ---------------- AUTHENTICATION ----------------
if not st.session_state.logged_in:
    st.title("🔒 NurseJoy.ai Authentication")

    option = st.selectbox("Choose an option", ["Login", "Sign Up"])
    roles = ["Doctor", "Patient", "Analyst"]

    if option == "Sign Up":
        st.subheader("Create a New Account")
        new_username = st.text_input("Choose a username")
        new_password = st.text_input("Choose a password", type="password")
        confirm_password = st.text_input("Confirm password", type="password")
        new_role = st.selectbox("Select your role", roles)

        if st.button("Sign Up"):
            if new_username.strip() == "" or new_password.strip() == "":
                st.warning("⚠️ Username and password cannot be empty.")
            elif new_username in users_db:
                st.error("❌ Username already exists. Please choose another.")
            elif new_password != confirm_password:
                st.error("❌ Passwords do not match.")
            else:
                hashed_pw = make_hashes(new_password)
                save_user(new_username, hashed_pw, new_role)
                users_db[new_username] = {"password": hashed_pw, "role": new_role}
                st.success("✅ Account created! Please login now.")
                st.stop()  # stop here to refresh UI

    elif option == "Login":
        st.subheader("Please log in to continue")
        username = st.text_input("Username")
        password = st.text_input("Password", type="password")
        selected_role = st.selectbox("Select your role", roles)

        if st.button("Login"):
            if username in users_db:
                user = users_db[username]
                if check_hashes(password, user["password"]) and selected_role == user["role"]:
                    st.session_state.logged_in = True
                    st.session_state.username = username
                    st.session_state.role = selected_role
                    st.success(f"✅ Welcome back, {username} ({selected_role})!")
                    st.stop()  # stop here to refresh UI
                else:
                    st.error("❌ Invalid username, password, or role.")
            else:
                st.error("❌ Invalid username, password, or role.")

    st.stop()

# ---------------- LOGOUT ----------------
st.sidebar.success(f"Logged in as: {st.session_state.username} ({st.session_state.role})")
if st.sidebar.button("Logout"):
    st.session_state.logged_in = False
    st.session_state.username = ""
    st.session_state.role = ""
    st.experimental_rerun()  # reload app after logout


#----------initialize jamai----------
# Initialize JamAI
jamai = JamAI(
    api_key="jamai_sk_71649a8af74bfb898c7300a4874b9ce29e088635870937ca",
    project_id="proj_267d212df1f8f66fdce751d0"
)

def extract_text_from_pdf(pdf_file):
    pdf = PdfReader(pdf_file)
    text = ""
    for page in pdf.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def generate_random_filename(extension=".docx"):
    random_str = ''.join(random.choices(string.ascii_letters + string.digits, k=10))
    return f"final_report_{random_str}{extension}"

# ---------------- MAIN APP ----------------
role = st.session_state.role

if role == "Doctor":
    if st.session_state.role == "Doctor":
        st.title("👩‍⚕️ Doctor Dashboard - Clinical Brief Generator")

        # Initialize JamAI once (if not done globally)
        jamai = JamAI(
            api_key="jamai_sk_71649a8af74bfb898c7300a4874b9ce29e088635870937ca",
            project_id="proj_267d212df1f8f66fdce751d0"
        )

        import random
        import string
        from PyPDF2 import PdfReader
        from io import BytesIO
        from docx import Document


        def extract_text_from_pdf(uploaded_file):
            reader = PdfReader(uploaded_file)
            return "\n".join(page.extract_text() for page in reader.pages if page.extract_text())


        def generate_random_filename():
            return ''.join(random.choices(string.ascii_letters + string.digits, k=8)) + ".docx"


        with st.container():
            st.header("🩺 Upload Patient Report for Clinical Brief")
            patient_pdf = st.file_uploader("📄 Upload Patient Report (PDF format)", type="pdf")

            if st.button("🚀 Generate Clinical Brief", use_container_width=True):
                if patient_pdf:
                    try:
                        report_text = extract_text_from_pdf(patient_pdf)

                        brief_completion = jamai.add_table_rows(
                            "action",
                            p.RowAddRequest(
                                table_id="ClinicalBriefGenerator",
                                data=[{"PatientReportRaw": report_text}],
                                stream=False
                            )
                        )

                        if brief_completion.rows:
                            row = brief_completion.rows[0].columns
                            critical_issues = row.get("CriticalIssues")
                            suggested_next_steps = row.get("SuggestedNextSteps")
                            secondary_concerns = row.get("SecondaryConcerns")
                            data_gaps = row.get("DataGaps")

                            st.subheader("✨ Generated Clinical Brief")
                            st.markdown(
                                f"""
                                <div class="generated-output">
                                    <h4>🚨 Critical Issues:</h4> <p>{critical_issues.text if critical_issues else 'N/A'}</p>
                                    <h4>🧭 Suggested Next Steps:</h4> <p>{suggested_next_steps.text if suggested_next_steps else 'N/A'}</p>
                                    <h4>📌 Secondary Concerns:</h4> <p>{secondary_concerns.text if secondary_concerns else 'N/A'}</p>
                                    <h4>🔍 Data Gaps:</h4> <p>{data_gaps.text if data_gaps else 'N/A'}</p>
                                </div>
                                """,
                                unsafe_allow_html=True
                            )

                            doc = Document()
                            doc.add_heading("Clinical Brief Report", level=1)
                            doc.add_heading("Critical Issues", level=2)
                            doc.add_paragraph(critical_issues.text if critical_issues else 'N/A')
                            doc.add_heading("Suggested Next Steps", level=2)
                            doc.add_paragraph(suggested_next_steps.text if suggested_next_steps else 'N/A')
                            doc.add_heading("Secondary Concerns", level=2)
                            doc.add_paragraph(secondary_concerns.text if secondary_concerns else 'N/A')
                            doc.add_heading("Data Gaps", level=2)
                            doc.add_paragraph(data_gaps.text if data_gaps else 'N/A')

                            buffer = BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            st.download_button(
                                label="📄 Download Clinical Brief as .docx",
                                data=buffer,
                                file_name=generate_random_filename(),
                                mime="application/octet-stream"
                            )
                        else:
                            st.error("⚠ Failed to process patient report.")
                    except Exception as e:
                        st.error(f"❌ Error during clinical brief generation: {e}")
                else:
                    st.warning("⚠ Please upload a patient report.")


elif role == "Patient":
    st.title("🌟 Malaysia's PokeNurse Centre - Gotta Save Em' All 🤾‍♀")


    # Safe text extractor helper
    def safe_text(value):
        return value.text if value and hasattr(value, "text") else "N/A"


    st.markdown("<h3>🩺 DoctorOak.ai - Patient Problem Diagnosis</h3>", unsafe_allow_html=True)

    with st.expander("📝 Enter Your Symptoms", expanded=True):
        patient_problem = st.text_area("Describe your health problem or symptoms here:")

        if st.button("🩺 Get Doctor Diagnosis", use_container_width=True):
            if patient_problem.strip():
                try:
                    # Call your JamAI API to add a row and get diagnosis output
                    booking_completion = jamai.add_table_rows(
                        "action",
                        p.RowAddRequest(
                            table_id="VirtualDoctor",
                            data=[{"Patient_Problem": patient_problem}],
                            stream=False
                        )
                    )

                    if booking_completion.rows:
                        row = booking_completion.rows[0].columns
                        doctor_diagnosis = safe_text(row.get("Doctor_Diagnosis"))

                        # Show diagnosis
                        st.subheader("💡 Doctor Diagnosis")
                        st.write(doctor_diagnosis)

                        # Generate downloadable Word report
                        doc = Document()
                        doc.add_heading("Virtual Doctor Diagnosis Report", level=1)

                        for label, value in [
                            ("Patient Problem", patient_problem),
                            ("Doctor Diagnosis", doctor_diagnosis),
                        ]:
                            para = doc.add_paragraph()
                            run_label = para.add_run(f"{label}: ")
                            run_label.bold = True
                            run_label.font.size = Pt(14)
                            run_value = para.add_run(str(value))
                            run_value.font.size = Pt(14)

                        buffer = BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)

                        st.download_button(
                            label="📄 Download Diagnosis Report (.docx)",
                            data=buffer,
                            file_name="VirtualDoctor_Diagnosis_Report.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        )
                    else:
                        st.error("⚠ No diagnosis returned from Virtual Doctor.")
                except Exception as e:
                    st.error(f"❌ Error retrieving diagnosis: {e}")
            else:
                st.warning("⚠ Please describe your health problem before submitting.")

    def safe_text(value):
        return value.text if value and hasattr(value, "text") else "N/A"

    st.markdown("<h3>🩺 NurseJoy.ai – Patient Support and Guidance Assistant</h3>", unsafe_allow_html=True)

    with st.expander("📅 Patient Booking Form with Brief Diagnosis", expanded=False):
        st.write("Please fill in your details below:")

        # --- Inputs ---
        full_name = st.text_input("👤 Full Name")
        dob = st.text_input("📅 Date of Birth (YYYY-MM-DD)")
        phone = st.text_input("📞 Handphone Number")
        email = st.text_input("📧 Email")
        symptoms = st.text_area("🤒 Symptoms")
        insurance = st.text_input("🏥 Insurance Provider")
        member_id = st.text_input("🆔 Member ID")
        accessibility = st.text_input("♿ Accessibility Requirements")
        language = st.text_input("🗣 Preferred Language")
        diet = st.text_input("🥗 Dietary Restrictions")

        # --- Submit ---
        if st.button("✅ Submit & Generate Booking", use_container_width=True):
            if full_name and dob and phone and email:
                try:
                    booking_completion = jamai.add_table_rows(
                        "action",
                        p.RowAddRequest(
                            table_id="BookingForm",  # Your JamAI table name
                            data=[{
                                "FullName": full_name,
                                "DOB": dob,
                                "HandphoneNumber": phone,
                                "Email": email,
                                "Symptoms": symptoms,
                                "InsuranceProvider": insurance,
                                "MemberID": member_id,
                                "Accessibility": accessibility,
                                "Language": language,
                                "DietRestriction": diet
                            }],
                            stream=False
                        )
                    )

                    if booking_completion.rows:
                        row = booking_completion.rows[0].columns

                        # --- Extract output fields safely ---
                        p_full_name = safe_text(row.get("PFullName"))
                        p_dob = safe_text(row.get("PDOB"))
                        p_phone = safe_text(row.get("PHandphoneNumber"))
                        p_email = safe_text(row.get("PEmail"))
                        p_insurance = safe_text(row.get("PInsuranceProvider"))
                        p_member_id = safe_text(row.get("PMemberID"))
                        p_accessibility = safe_text(row.get("PAccessibility"))
                        p_language = safe_text(row.get("PLanguage"))
                        p_diet = safe_text(row.get("PDietRestriction"))
                        possible_illness = safe_text(row.get("PossibleIllness"))
                        possible_treatment = safe_text(row.get("PossibleTreatment"))
                        appt_date = safe_text(row.get("ApptDate"))
                        appt_location = safe_text(row.get("ApptLocation"))
                        dept = safe_text(row.get("SpecialistDepartment"))
                        doctor = safe_text(row.get("DoctorName"))
                        booking_id = safe_text(row.get("BookingID"))

                        # --- Show results in UI ---
                        st.subheader("📋 Booking Confirmation")
                        st.markdown(f"""
                            Full Name: {p_full_name}  
                            DOB: {p_dob}  
                            Phone: {p_phone}  
                            Email: {p_email}  
                            Insurance Provider: {p_insurance}  
                            Member ID: {p_member_id}  
                            Accessibility: {p_accessibility}  
                            Language: {p_language}  
                            Diet Restriction: {p_diet}  
                            Possible Illness: {possible_illness}  
                            Possible Treatment: {possible_treatment}  
                            Appointment Date: {appt_date}  
                            Appointment Location: {appt_location}  
                            Specialist Department: {dept}  
                            Doctor Name: {doctor}  
                            Booking ID: {booking_id}  
                        """)

                        # --- Create Word report ---
                        doc = Document()
                        doc.add_heading("Patient Booking Confirmation", level=1)

                        # Large readable font for elderly
                        for label, value in [
                            ("Full Name", p_full_name),
                            ("DOB", p_dob),
                            ("Phone", p_phone),
                            ("Email", p_email),
                            ("Insurance Provider", p_insurance),
                            ("Member ID", p_member_id),
                            ("Accessibility", p_accessibility),
                            ("Language", p_language),
                            ("Diet Restriction", p_diet),
                            ("Possible Illness", possible_illness),
                            ("Possible Treatment", possible_treatment),
                            ("Appointment Date", appt_date),
                            ("Appointment Location", appt_location),
                            ("Specialist Department", dept),
                            ("Doctor Name", doctor),
                            ("Booking ID", booking_id),
                        ]:
                            para = doc.add_paragraph()
                            run_label = para.add_run(f"{label}: ")
                            run_label.bold = True
                            run_label.font.size = Pt(14)
                            run_value = para.add_run(str(value))
                            run_value.font.size = Pt(14)

                        buffer = BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)

                        # --- Download button ---
                        st.download_button(
                            label="📄 Download Booking Report (.docx)",
                            data=buffer,
                            file_name=f"Booking_{booking_id}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    else:
                        st.error("⚠ No booking details were returned.")
                except Exception as e:
                    st.error(f"❌ Error submitting booking: {e}")
            else:
                st.warning("⚠ Please fill in all required fields (Name, DOB, Phone, Email).")


    # --- Medicine Identifier UI ---
    with st.expander("💊 Medicine Identifier", expanded=False):
        st.write('Identify your medicine and learn how to use it safely ✦˖ ࣪𖤐₊˚ೄ')

        uploaded_file = st.file_uploader("📷 Upload a clear image of your medicine", type=["jpg", "jpeg", "png"])
        analyze_button = st.button("🔍 Identify Medicine", use_container_width=True)

        if uploaded_file and analyze_button:
            st.write("Filename:", uploaded_file.name)

            # Convert uploaded file to JPEG
            with NamedTemporaryFile(delete=False, suffix=".jpeg") as temp_file:
                try:
                    image = Image.open(uploaded_file).convert("RGB")
                    image.save(temp_file, format="JPEG")
                    temp_file_path = temp_file.name
                except Exception as e:
                    st.error(f"❌ Image processing error: {e}")
                    st.stop()

            # Upload image to JamAI
            try:
                upload_response = jamai.file.upload_file(temp_file_path)
            except Exception as e:
                st.error(f"❌ Failed to upload image: {e}")
                st.stop()

            # Identify the medicine
            try:
                completion = jamai.table.add_table_rows(
                    "action",
                    p.RowAddRequest(
                        table_id="MedicineIdentifier",
                        data=[dict(Medicine=upload_response.uri)],
                        stream=False,
                    ),
                )
            except Exception as e:
                st.error(f"❌ Medicine identification request failed: {e}")
                st.stop()

            # Display image preview
            image_bytes = uploaded_file.getvalue()
            encoded_image = base64.b64encode(image_bytes).decode()
            st.image(f"data:image/jpeg;base64,{encoded_image}", width=300, caption="Uploaded Medicine Image")

            if completion.rows:
                st.success("Medicine successfully identified!")
            else:
                st.error("Identification failed. Please try again.")

            # Fetch result from the table
            rows = jamai.table.list_table_rows("action", "MedicineIdentifier")
            if rows.items:
                row = rows.items[0]

                medicine_attributes = {
                    "Medicine Name": row.get("MedicineName", {}).get("value", "N/A"),
                    "Information": row.get("MedicineInformation", {}).get("value", "N/A"),
                    "Dosage & Usage": row.get("MedicineDosageUsage", {}).get("value", "N/A"),
                }

                st.subheader("📋 Medicine Details")
                for key, value in medicine_attributes.items():
                    st.markdown(f"{key}: {value}")

                # Downloadable .docx report
                with st.container():
                    st.subheader("📥 Download Medicine Report")
                    doc = Document()
                    doc.add_heading("Medicine Information Report", level=1)

                    # Set big readable font for elderly
                    style = doc.styles['Normal']
                    font = style.font
                    font.size = Pt(14)

                    for key, value in medicine_attributes.items():
                        doc.add_heading(key, level=2)
                        doc.add_paragraph(value)

                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)

                    st.download_button(
                        label="📄 Download Medicine Report as .docx",
                        data=buffer,
                        file_name="Medicine_Report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            else:
                st.warning("⚠ No results yet. Please try again after a moment.")

    # --- Nutrition Menu UI ---
    with st.expander("🥗 Nutrition Menu Generator", expanded=False):
        st.write('Healthy recipes based on your fridge and condition ✦˖ ࣪𖤐₊˚ೄ')

        # Inputs
        ingredients_text = st.text_area("📝 List your refrigerator ingredients (optional if you upload image)")
        uploaded_image = st.file_uploader("📷 Upload a photo of your refrigerator contents", type=["jpg", "jpeg", "png"])
        illness_text = st.text_input("💊 Enter your illness or health condition")

        generate_button = st.button("🍳 Generate Healthy Menu", use_container_width=True)

        if generate_button:
            if not (ingredients_text or uploaded_image):
                st.warning("⚠ Please enter ingredients text or upload a fridge image.")
                st.stop()

            if not illness_text:
                st.warning("⚠ Please enter your illness or health condition.")
                st.stop()

            # Handle image upload
            fridge_image_uri = None
            if uploaded_image:
                with NamedTemporaryFile(delete=False, suffix=".jpeg") as temp_file:
                    try:
                        image = Image.open(uploaded_image).convert("RGB")
                        image.save(temp_file, format="JPEG")
                        temp_file_path = temp_file.name
                    except Exception as e:
                        st.error(f"❌ Image processing error: {e}")
                        st.stop()

                try:
                    upload_response = jamai.file.upload_file(temp_file_path)
                    fridge_image_uri = upload_response.uri
                except Exception as e:
                    st.error(f"❌ Failed to upload image: {e}")
                    st.stop()

                # Display image preview
                image_bytes = uploaded_image.getvalue()
                encoded_image = base64.b64encode(image_bytes).decode()
                st.image(f"data:image/jpeg;base64,{encoded_image}", width=300, caption="Uploaded Refrigerator Image")

            # Send request to NutritionMenu table
            try:
                completion = jamai.table.add_table_rows(
                    "action",
                    p.RowAddRequest(
                        table_id="NutritionMenu",
                        data=[{
                            "RefrigeratorIngredients": ingredients_text if ingredients_text else "",
                            "RefrigeratorIngredientsImage": fridge_image_uri if fridge_image_uri else "",
                            "Illness": illness_text
                        }],
                        stream=False,
                    ),
                )
            except Exception as e:
                st.error(f"❌ Menu generation request failed: {e}")
                st.stop()

            if completion.rows:
                st.success("✅ Menu recipe generated!")
            else:
                st.error("⚠ Failed to generate menu. Please try again.")

            # Fetch result
            rows = jamai.table.list_table_rows("action", "NutritionMenu")
            if rows.items:
                row = rows.items[0]
                menu_recipe = row.get("MenuRecipe", {}).get("value", "N/A")

                st.subheader("🍽 Suggested Menu Recipe")
                st.markdown(menu_recipe)

                # Downloadable docx
                with st.container():
                    st.subheader("📥 Download Menu Recipe Report")
                    doc = Document()
                    doc.add_heading("Nutrition Menu Recipe", level=1)

                    # Large font for elderly readability
                    style = doc.styles['Normal']
                    font = style.font
                    font.size = Pt(14)

                    doc.add_heading("Illness / Health Condition", level=2)
                    doc.add_paragraph(illness_text)

                    doc.add_heading("Ingredients Provided", level=2)
                    doc.add_paragraph(ingredients_text if ingredients_text else "Uploaded fridge image only.")

                    doc.add_heading("Suggested Menu Recipe", level=2)
                    doc.add_paragraph(menu_recipe)

                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)

                    st.download_button(
                        label="📄 Download Recipe as .docx",
                        data=buffer,
                        file_name="NutritionMenu_Report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            else:
                st.warning("⚠ No results yet. Please try again after a moment.")

if st.button("🗺 Indoor Mapping Navigator"):
    js = "window.open('https://www.hospitalseribotani.com.my/360-virtual-tour/', '_blank')"  # New tab
    st.components.v1.html(f"<script>{js}</script>")

if role == "Analyst":
    if st.session_state.role == "Analyst":
     st.title("👩‍⚕️ Analyst Dashboard")


     # --- Safe text helper ---
     def safe_text(value):
         return value.text if value and hasattr(value, "text") else "N/A"


     st.title("🛡 Health Signal Summarizer")
     st.write("Paste public health-related reports or social media/news snippets below:")

     # --- Inputs ---
     raw_reports = st.text_area("📰 Raw Reports (paste text from social media, news, or community tips)")
     report_source = st.text_input("🌐 Report Source(s) (e.g., Twitter, Facebook, Local News)")
     report_date = st.text_input("📅 Report Date/Time (YYYY-MM-DD HH:MM)")
     reporter_name = st.text_input("👤 Reporter Name (optional)")
     location_mentioned = st.text_input("📍 Location Mentioned")

     # --- Submit ---
     if st.button("✅ Analyze Health Signals", use_container_width=True):
         if raw_reports and report_source and report_date:
             try:
                 summarization = jamai.add_table_rows(
                     "action",
                     p.RowAddRequest(
                         table_id="TBL_HealthSignal",  # Your Jamaibase table name
                         data=[{
                             "RawReports": raw_reports,
                             "ReportSource": report_source,
                             "ReportDate": report_date,
                             "ReporterName": reporter_name,
                             "LocationMentioned": location_mentioned
                         }],
                         stream=False
                     )
                 )

                 if summarization.rows:
                     row = summarization.rows[0].columns

                     # --- Extract output fields safely ---
                     alert_summary = safe_text(row.get("AlertSummary"))
                     location = safe_text(row.get("Location"))
                     threat_type = safe_text(row.get("PossibleThreatType"))
                     credibility = safe_text(row.get("CredibilityScore"))
                     trend = safe_text(row.get("TrendIndicator"))
                     urgency = safe_text(row.get("UrgencyFlag"))

                     # --- Show results in UI ---
                     st.subheader("📋 Health Signal Summary")
                     st.markdown(f"""
                         Alert Summary: {alert_summary}  
                         Location: {location}  
                         Threat Type: {threat_type}  
                         Credibility: {credibility}  
                         Trend: {trend}  
                         Urgency: {urgency}  
                     """)

                     # --- Create Word report ---
                     doc = Document()
                     doc.add_heading("Health Signal Report", level=1)

                     for label, value in [
                         ("Alert Summary", alert_summary),
                         ("Location", location),
                         ("Threat Type", threat_type),
                         ("Credibility", credibility),
                         ("Trend", trend),
                         ("Urgency", urgency),
                         ("Source", report_source),
                         ("Reported Date/Time", report_date),
                         ("Reporter Name", reporter_name),
                         ("Location Mentioned", location_mentioned),
                         ("Raw Reports", raw_reports),
                     ]:
                         para = doc.add_paragraph()
                         run_label = para.add_run(f"{label}: ")
                         run_label.bold = True
                         run_label.font.size = Pt(14)
                         run_value = para.add_run(str(value))
                         run_value.font.size = Pt(14)

                     buffer = BytesIO()
                     doc.save(buffer)
                     buffer.seek(0)

                     st.download_button(
                         label="📄 Download Health Signal Report (.docx)",
                         data=buffer,
                         file_name=f"HealthSignal_{location}.docx",
                         mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                     )
                 else:
                     st.error("⚠ No health signal data returned.")
             except Exception as e:
                 st.error(f"❌ Error analyzing health signals: {e}")
         else:
             st.warning("⚠ Please fill in required fields: Raw Reports, Report Source, and Report Date/Time.")

